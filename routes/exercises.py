"""Exercise routes — serve pre-generated exercises with answer reveal + Word download + error tracking."""

import html as _html
import json
import io
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, Depends, Request, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor

from db.database import get_db
from db.models import ErrorLog
from routes.pages import MATH_SECTIONS, get_nav_subjects, NAMES
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.ai_tutor import AITutorService
from services.progress import ProgressService
from config import settings

router = APIRouter()

EXERCISE_CACHE_DIR = Path("data/exercises")
EXERCISE_CACHE_DIR.mkdir(parents=True, exist_ok=True)

# Cache loaded exercises by subject
_EXERCISE_CACHE: dict[str, dict] = {}

def _load_exercises(subject: str = "math") -> dict:
    """Load exercises from per-subject JSON file."""
    if subject in _EXERCISE_CACHE:
        return _EXERCISE_CACHE[subject]
    json_path = Path(f"data/exercises/{subject}.json")
    if json_path.exists():
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            _EXERCISE_CACHE[subject] = data
            return data
    # Fallback: old math exercises.json
    fallback = Path("data/exercises/exercises.json")
    if fallback.exists():
        with open(fallback, "r", encoding="utf-8") as f:
            data = json.load(f)
            _EXERCISE_CACHE[subject] = data
            return data
    return _DEFAULT_EXERCISES

_DEFAULT_EXERCISES = {
    "default": [
        {"type": "choice", "question": "请先学习教材内容，再来做练习哦！", "choices": ["知道了"], "answer": 0, "explanation": "这个知识点的练习题正在准备中，先仔细阅读教材吧！"},
    ],
}

# Legacy hardcoded exercises (for sections without JSON data)
_LEGACY_EXERCISES = {
    "ch01_sec01": [
        {"type": "choice", "question": "下列哪个是负数？", "choices": ["5", "0", "-3", "1/2"], "answer": 2, "explanation": "负数是小于0的数，-3是负数。"},
        {"type": "choice", "question": "如果上升5m记作+5m，那么下降3m记作什么？", "choices": ["+3m", "-3m", "0m", "3m"], "answer": 1, "explanation": "具有相反意义的量，下降用负数表示，所以是-3m。"},
        {"type": "choice", "question": "0是正数还是负数？", "choices": ["正数", "负数", "既是正数也是负数", "既不是正数也不是负数"], "answer": 3, "explanation": "0既不是正数也不是负数，它是正数和负数的分界点。"},
        {"type": "fill", "question": "比0大5的数是___，比0小3的数是___。", "answer": "+5（或5）, -3", "explanation": "比0大的数是正数，比0小的数是负数。"},
        {"type": "choice", "question": "如果向东走50m记作+50m，那么-30m表示什么？", "choices": ["向东走30m", "向西走30m", "向南走30m", "向北走30m"], "answer": 1, "explanation": "正数表示向东，那么负数表示相反方向即向西。"},
    ],
    "ch01_sec02": [
        {"type": "choice", "question": "下列哪个是有理数？", "choices": ["π", "√2", "-5", "e"], "answer": 2, "explanation": "有理数是可以写成分数形式的数，-5=-5/1是有理数。"},
        {"type": "choice", "question": "在数轴上，-3和2哪个大？", "choices": ["-3大", "2大", "一样大", "无法比较"], "answer": 1, "explanation": "在数轴上右边的数总比左边的大，2在-3的右边。"},
        {"type": "fill", "question": "|-5|的值是___。", "answer": "5", "explanation": "绝对值表示数轴上的点到原点的距离，|-5|=5。"},
        {"type": "choice", "question": "下列各数中，最小的是？", "choices": ["-2", "0", "1", "-5"], "answer": 3, "explanation": "负数中，绝对值大的反而小，|-5|>|-2|，所以-5最小。"},
        {"type": "fill", "question": "比-3大比2小的整数有___个。", "answer": "4（-2, -1, 0, 1）", "explanation": "-2, -1, 0, 1都在-3和2之间。"},
    ],
    "ch02_sec01": [
        {"type": "choice", "question": "计算(-3) + (-5)的结果是？", "choices": ["2", "-2", "-8", "8"], "answer": 2, "explanation": "同号相加：符号不变，绝对值相加。(-3)+(-5)=-(3+5)=-8。"},
        {"type": "choice", "question": "计算(-7) + 4的结果是？", "choices": ["-11", "-3", "3", "11"], "answer": 1, "explanation": "异号相加：取绝对值大的符号，用较大绝对值减较小绝对值。|-7|>|4|，所以7-4=3，符号为负，结果是-3。"},
        {"type": "fill", "question": "5 + (-5) = ___。", "answer": "0", "explanation": "互为相反数的两个数相加得0。"},
        {"type": "choice", "question": "把(-8) - (-3)转化为加法是？", "choices": ["(-8) + 3", "(-8) + (-3)", "8 + (-3)", "8 + 3"], "answer": 0, "explanation": "减去一个数等于加上它的相反数：(-8)-(-3)=(-8)+3=-5。"},
        {"type": "fill", "question": "冰箱冷藏室温度3℃，冷冻室温度-18℃，温差是___℃。", "answer": "21", "explanation": "温差=3-(-18)=3+18=21℃。"},
    ],
    "ch02_sec02": [
        {"type": "choice", "question": "计算(-3) × 4的结果是？", "choices": ["12", "-12", "7", "-7"], "answer": 1, "explanation": "异号得负，绝对值相乘。(-3)×4=-(3×4)=-12。"},
        {"type": "choice", "question": "计算(-2) × (-5)的结果是？", "choices": ["-10", "-7", "7", "10"], "answer": 3, "explanation": "同号得正，绝对值相乘。(-2)×(-5)=+(2×5)=10。"},
        {"type": "fill", "question": "-1/2的倒数是___。", "answer": "-2", "explanation": "乘积为1的两个数互为倒数，(-1/2)×(-2)=1。"},
        {"type": "fill", "question": "计算(-12) ÷ (-3) = ___。", "answer": "4", "explanation": "除以一个数等于乘它的倒数：(-12)÷(-3)=(-12)×(-1/3)=4。"},
        {"type": "choice", "question": "用分配律计算(-8)×(25-5)，以下哪个正确？", "choices": ["(-8)×20=-160", "(-8)×25-(-8)×5=-200+40=-160", "(-8)×25+(-8)×5=-200-40=-240", "8×25-8×5=160"], "answer": 1, "explanation": "分配律：a(b+c)=ab+ac，所以(-8)×25+(-8)×(-5)=-200+40=-160。"},
    ],
    "ch02_sec03": [
        {"type": "choice", "question": "(-2)³的值是？", "choices": ["-8", "8", "-6", "6"], "answer": 0, "explanation": "负数的奇次幂是负数。(-2)³=(-2)×(-2)×(-2)=-8。"},
        {"type": "choice", "question": "(-3)²的值是？", "choices": ["-9", "9", "-6", "6"], "answer": 1, "explanation": "负数的偶次幂是正数。(-3)²=(-3)×(-3)=9。"},
        {"type": "fill", "question": "-2² = ___。", "answer": "-4", "explanation": "注意：-2²=-(2×2)=-4，而(-2)²=4，括号很重要！"},
        {"type": "fill", "question": "用科学记数法表示5380000 = ___。", "answer": "5.38×10⁶", "explanation": "科学记数法：a×10ⁿ，其中1≤|a|<10。5380000=5.38×10⁶。"},
        {"type": "choice", "question": "计算2×(-3)²-(-4)的结果是？", "choices": ["14", "22", "-14", "10"], "answer": 1, "explanation": "先乘方：(-3)²=9，再乘：2×9=18，最后减：18-(-4)=18+4=22。"},
    ],
    "ch03_sec01": [
        {"type": "choice", "question": "下列哪个是代数式？", "choices": ["3+2=5", "x+3", "x>5", "S=ab"], "answer": 1, "explanation": "代数式是用运算符号连接的式子，不含等号或不等号。x+3是代数式。"},
        {"type": "fill", "question": "\"a的3倍与b的一半的和\"用代数式表示为___。", "answer": "3a + b/2", "explanation": "a的3倍=3a，b的一半=b/2，和=3a+b/2。"},
        {"type": "choice", "question": "长方形长a宽b，周长用代数式表示为？", "choices": ["ab", "2a+b", "2(a+b)", "a²+b²"], "answer": 2, "explanation": "长方形周长=2×(长+宽)=2(a+b)。"},
        {"type": "fill", "question": "温度由t℃下降5℃后是___℃。", "answer": "t-5", "explanation": "下降5℃就是减5，所以是t-5。"},
        {"type": "choice", "question": "代数式3x+2y表示什么？", "choices": ["x的3倍加y的2倍", "3加x加2加y", "3x+2再乘y", "x加y的3倍加2"], "answer": 0, "explanation": "3x表示x的3倍，2y表示y的2倍，相加表示它们的和。"},
    ],
    "ch03_sec02": [
        {"type": "fill", "question": "当a=2时，代数式3a+1的值是___。", "answer": "7", "explanation": "代入a=2：3×2+1=6+1=7。"},
        {"type": "fill", "question": "当x=-3时，代数式x²+2x的值是___。", "answer": "3", "explanation": "代入x=-3：(-3)²+2×(-3)=9-6=3。注意负数要加括号！"},
        {"type": "choice", "question": "当a=-1时，-a²的值是？", "choices": ["1", "-1", "0", "2"], "answer": 1, "explanation": "-a²=-(a²)=-((-1)²)=-1。"},
        {"type": "fill", "question": "若|x|=3，则x=___。", "answer": "3或-3", "explanation": "绝对值等于3的数有两个：3和-3。"},
        {"type": "choice", "question": "当n=4时，代数式n(n+1)/2的值是？", "choices": ["6", "8", "10", "12"], "answer": 2, "explanation": "代入n=4：4×(4+1)÷2=4×5÷2=20÷2=10。"},
    ],
    "ch04_sec01": [
        {"type": "choice", "question": "单项式-3x²y的系数是？", "choices": ["3", "-3", "2", "-1"], "answer": 1, "explanation": "系数是数字因数（含符号），-3x²y的系数是-3。"},
        {"type": "fill", "question": "单项式5ab³c的次数是___次。", "answer": "5", "explanation": "所有字母指数之和：a¹+b³+c¹=1+3+1=5次。"},
        {"type": "choice", "question": "多项式3x²-2x+1的项数是？", "choices": ["1项", "2项", "3项", "4项"], "answer": 2, "explanation": "多项式3x²-2x+1有三项：3x²、-2x、1。"},
        {"type": "fill", "question": "多项式4x³-2x²+x-7中，次数最高的项是___。", "answer": "4x³（3次）", "explanation": "各项次数：4x³是3次，-2x²是2次，x是1次，-7是0次。最高次数是3次。"},
        {"type": "choice", "question": "下列哪个是单项式？", "choices": ["x+y", "3ab", "x²-1", "a/b+1"], "answer": 1, "explanation": "单项式是数字与字母的乘积。3ab是单项式，其他都是多项式或分式。"},
    ],
    "ch04_sec02": [
        {"type": "choice", "question": "3x²y和-5x²y是同类项吗？", "choices": ["是", "不是", "有时是", "无法判断"], "answer": 0, "explanation": "同类项要求字母相同且相同字母的指数相同。两者都是x²y，只是系数不同，是同类项。"},
        {"type": "fill", "question": "5a+3a = ___。", "answer": "8a", "explanation": "合并同类项：系数相加5+3=8，字母部分a不变。"},
        {"type": "fill", "question": "化简：-(a-b) = ___。", "answer": "-a+b", "explanation": "括号前是\"-\"，去括号时各项都变号：-(a-b)=-a+b。"},
        {"type": "choice", "question": "化简(2x²+3x-1)+(x²-2x+4)的结果是？", "choices": ["3x²+x+3", "x²+5x+3", "3x²+5x+3", "3x²+x-5"], "answer": 0, "explanation": "去括号：2x²+3x-1+x²-2x+4，合并同类项：(2+1)x²+(3-2)x+(-1+4)=3x²+x+3。"},
        {"type": "fill", "question": "化简：5x-(2x-3) = ___。", "answer": "3x+3", "explanation": "5x-(2x-3)=5x-2x+3=3x+3。注意：-(2x-3)去括号后-2x变+3。"},
    ],
    "ch05_sec01": [
        {"type": "choice", "question": "下列哪个是一元一次方程？", "choices": ["x²+2=6", "2x+3=7", "x+y=5", "2x>8"], "answer": 1, "explanation": "一元一次方程：一个未知数，次数为1，是等式。2x+3=7符合。"},
        {"type": "fill", "question": "方程3x=12的解是x=___。", "answer": "4", "explanation": "两边同除以3：x=12÷3=4。"},
        {"type": "fill", "question": "若x=2是方程2x+a=10的解，则a=___。", "answer": "6", "explanation": "代入x=2：2×2+a=10，4+a=10，a=6。"},
        {"type": "choice", "question": "方程x+5=3的解是？", "choices": ["x=8", "x=2", "x=-2", "x=-8"], "answer": 2, "explanation": "x+5=3，两边减5：x=3-5=-2。"},
        {"type": "choice", "question": "下列变形正确的是？", "choices": ["若a=b，则a+c=b-c", "若a=b，则ac=bc", "若ac=bc，则a=b", "若a²=b²，则a=b"], "answer": 1, "explanation": "等式性质2：若a=b，则ac=bc（两边同乘一个数，等式仍成立）。"},
    ],
    "ch05_sec02": [
        {"type": "fill", "question": "解方程：2x+3=11，x=___。", "answer": "4", "explanation": "移项：2x=11-3=8，系数化为1：x=4。"},
        {"type": "fill", "question": "解方程：3(x-2)=12，x=___。", "answer": "6", "explanation": "去括号：3x-6=12，移项：3x=18，x=6。"},
        {"type": "choice", "question": "解方程x/2 - 1 = 3，x=？", "choices": ["4", "6", "8", "2"], "answer": 2, "explanation": "x/2-1=3，移项：x/2=4，两边乘2：x=8。"},
        {"type": "fill", "question": "方程(2x-1)/3 = 5的解是x=___。", "answer": "8", "explanation": "两边乘3：2x-1=15，移项：2x=16，x=8。"},
        {"type": "choice", "question": "解方程5x-7=3x+5，第一步应该？", "choices": ["合并同类项", "移项", "去括号", "系数化为1"], "answer": 1, "explanation": "先移项将含x的项移到一边：5x-3x=5+7，再合并：2x=12，x=6。"},
    ],
    "ch05_sec03": [
        {"type": "fill", "question": "小明买了3支笔和2个本子，共花了17元。若每支笔x元，每本y元，可列方程___。", "answer": "3x+2y=17", "explanation": "总花费=笔的钱+本子的钱=3x+2y=17。"},
        {"type": "choice", "question": "甲乙两地相距120km，汽车从甲地到乙地用了2小时，求速度。设速度为x km/h，方程是？", "choices": ["x+2=120", "2x=120", "x/2=120", "x-2=120"], "answer": 1, "explanation": "路程=速度×时间，120=2x。"},
        {"type": "fill", "question": "一个数的3倍加5等于这个数的2倍减1，设这个数为x，可列方程___。", "answer": "3x+5=2x-1，解得x=-6", "explanation": "3x+5=2x-1，移项：3x-2x=-1-5，x=-6。"},
        {"type": "choice", "question": "某商品进价100元，按20%的利润率定价，售价是多少？", "choices": ["120元", "102元", "80元", "20元"], "answer": 0, "explanation": "售价=进价×(1+利润率)=100×(1+20%)=120元。"},
        {"type": "fill", "question": "一个两位数，十位数字是a，个位数字是b，这个两位数表示为___。", "answer": "10a+b", "explanation": "十位数字a表示10a，个位数字b表示b，两位数=10a+b。"},
    ],
    "ch06_sec01": [
        {"type": "choice", "question": "下列哪个是锥体？", "choices": ["正方体", "圆柱", "圆锥", "球"], "answer": 2, "explanation": "锥体有一个顶点和一个底面，侧面是曲面。圆锥是锥体。"},
        {"type": "choice", "question": "正方体有几个面？", "choices": ["4个", "6个", "8个", "12个"], "answer": 1, "explanation": "正方体有6个面（上、下、左、右、前、后），都是正方形。"},
        {"type": "fill", "question": "圆柱的侧面展开图是___形。", "answer": "长方", "explanation": "圆柱的侧面展开后是一个长方形。"},
        {"type": "fill", "question": "圆锥的侧面展开图是___形。", "answer": "扇", "explanation": "圆锥的侧面展开后是一个扇形。"},
        {"type": "choice", "question": "从正面看一个圆柱，看到的图形是？", "choices": ["圆", "长方形", "三角形", "梯形"], "answer": 1, "explanation": "从正面（侧面）看圆柱，看到的是一个长方形。"},
    ],
    "ch06_sec02": [
        {"type": "choice", "question": "经过两点可以画几条直线？", "choices": ["0条", "1条", "2条", "无数条"], "answer": 1, "explanation": "两点确定一条直线，经过两点只能画一条直线。"},
        {"type": "fill", "question": "线段AB=10cm，C是AB的中点，则AC=___cm。", "answer": "5", "explanation": "中点将线段分成两等份，AC=AB÷2=10÷2=5cm。"},
        {"type": "choice", "question": "下列说法正确的是？", "choices": ["直线有两个端点", "射线有一个端点", "线段没有端点", "射线有两个端点"], "answer": 1, "explanation": "直线0个端点，射线1个端点，线段2个端点。"},
        {"type": "fill", "question": "从A地到B地有三条路，最短的是___。", "answer": "线段AB（直路）", "explanation": "两点之间线段最短。"},
        {"type": "choice", "question": "下列哪个工具用来测量两点之间的距离？", "choices": ["量角器", "三角板", "刻度尺", "圆规"], "answer": 2, "explanation": "测量线段长度用刻度尺。"},
    ],
    "ch06_sec03": [
        {"type": "fill", "question": "35° = ___′（分）。", "answer": "2100", "explanation": "1°=60′，35×60=2100′。"},
        {"type": "choice", "question": "55°的余角是多少度？", "choices": ["35°", "45°", "125°", "55°"], "answer": 0, "explanation": "互余两角之和为90°，90°-55°=35°。"},
        {"type": "fill", "question": "120°的补角是___°。", "answer": "60", "explanation": "互补两角之和为180°，180°-120°=60°。"},
        {"type": "choice", "question": "下列哪个角是钝角？", "choices": ["30°", "90°", "120°", "180°"], "answer": 2, "explanation": "钝角：大于90°且小于180°，120°是钝角。"},
        {"type": "choice", "question": "OC是∠AOB的平分线，若∠AOB=80°，则∠AOC=？", "choices": ["40°", "80°", "160°", "20°"], "answer": 0, "explanation": "角平分线将角分成相等的两部分，80°÷2=40°。"},
    ],
    "ch01_reading": [
        {"type": "fill", "question": "某零件标准尺寸为50mm，允许偏差±0.02mm，合格范围是___~___mm。", "answer": "49.98~50.02", "explanation": "最小=50-0.02=49.98mm，最大=50+0.02=50.02mm。"},
        {"type": "choice", "question": "偏差±0.05mm的含义是？", "choices": ["只能是正偏差", "只能是负偏差", "可正可负，绝对值不超过0.05", "偏差必须为0"], "answer": 2, "explanation": "±0.05表示允许比标准值大0.05或小0.05。"},
        {"type": "fill", "question": "某食品标注净含量200g±5g，最少含量为___g。", "answer": "195", "explanation": "最少=200-5=195g。"},
        {"type": "choice", "question": "为什么工业生产中需要允许偏差？", "choices": ["工人偷懒", "节省材料", "完全精确不可能也不必要", "故意做不准"], "answer": 2, "explanation": "实际加工无法做到绝对精确，在一定范围内不影响使用即可。"},
        {"type": "fill", "question": "某产品标准质量m克，偏差d克，合格品的质量范围是___。", "answer": "m-d ≤ 实际质量 ≤ m+d", "explanation": "标准值±偏差定义了合格范围。"},
    ],
    "ch01_history": [
        {"type": "choice", "question": "最早提出正负数概念的是哪本数学著作？", "choices": ["《几何原本》", "《九章算术》", "《周髀算经》", "《孙子算经》"], "answer": 1, "explanation": "《九章算术》约成书于1世纪，在\"方程\"章中首次提出正负数概念及运算法则。"},
        {"type": "fill", "question": "中国古代用___色算筹表示正数，___色算筹表示负数。", "answer": "红、黑", "explanation": "红色为正，黑色为负，这是世界上最早的负数表示方法之一。"},
        {"type": "choice", "question": "负数从被提出到被欧洲数学家广泛接受，大约经历了多长时间？", "choices": ["几十年", "一百年", "上千年", "从未被接受"], "answer": 2, "explanation": "中国在1世纪就使用负数，但欧洲直到17世纪才逐渐接受，经历了上千年。"},
        {"type": "fill", "question": "《九章算术》中\"正与负\"用来表示___与___。", "answer": "卖出与买入（或收入与支出）", "explanation": "《九章算术》用\"正与负\"表示买卖中的钱数，卖出为正，买入为负。"},
        {"type": "choice", "question": "这段数学史给我们的启示是？", "choices": ["数学概念都是理所当然的", "数学概念的发展需要漫长时间", "负数没有实际用途", "古人不需要负数"], "answer": 1, "explanation": "数学概念从提出到被广泛接受，往往需要漫长的历史过程。"},
    ],
    "default": [
        {"type": "choice", "question": "请先学习教材内容，再来做练习哦！", "choices": ["知道了"], "answer": 0, "explanation": "这个知识点的练习题正在准备中，先仔细阅读教材吧！"},
    ],
}

# Backward compat: default to math exercises
EXERCISES = _load_exercises("math")

def get_exercises_for(subject: str = "math") -> dict:
    """Get exercises for a specific subject (loaded from JSON or legacy fallback)."""
    data = _load_exercises(subject)
    # If only has "default", try legacy math data for math subject
    if subject == "math" and len(data) <= 2:
        # Merge legacy math exercises
        for k, v in _LEGACY_EXERCISES.items():
            if k not in data:
                data[k] = v
        if "default" not in data:
            data["default"] = _DEFAULT_EXERCISES["default"]
    return data


def _build_exercise_html(section_id: str, subject: str = "math") -> str:
    """Build HTML for exercises using ae-card style."""
    ex_data = get_exercises_for(subject)
    ex_list = ex_data.get(section_id, ex_data.get("default", []))
    if not ex_list:
        ex_list = EXERCISES["default"]

    parts = []
    for i, ex in enumerate(ex_list):
        qid = f"ae-{section_id}-{i}"
        qtype_icon = {"choice": "🔤", "fill": "✏️", "true_false": "✅"}.get(ex["type"], "📝")
        parts.append(f'<div class="ae-card" id="{qid}">')
        parts.append(
            f'<div class="ae-q">'
            f'<span class="ae-q-num">{i+1}</span>'
            f'<span class="ae-q-type">{qtype_icon}</span>'
            f'<span class="ae-q-text">{ex["question"]}</span>'
            f'</div>'
        )

        if ex["type"] == "choice" and "choices" in ex:
            parts.append('<div class="ae-choices">')
            for j, choice in enumerate(ex.get("choices", [])):
                parts.append(
                    f'<button class="ae-opt" onclick="aeCheck(\'{qid}\',{j},{ex["answer"]},this)" data-correct="{ex["answer"]}">'
                    f'<span class="ae-opt-letter">{chr(65+j)}</span>'
                    f'<span class="ae-opt-text">{choice}</span>'
                    f'</button>'
                )
            parts.append('</div>')
        elif ex["type"] == "fill":
            parts.append(
                f'<div class="ae-fill">'
                f'<input class="ae-input" id="{qid}-inp" placeholder="输入你的答案...">'
                f'<button class="ae-btn" onclick="aeFill(\'{qid}\',\'{ex["answer"]}\')">检查</button>'
                f'</div>'
            )
        elif ex["type"] == "true_false":
            correct_idx = 0 if str(ex["answer"]) == "正确" else 1
            parts.append(
                f'<div class="ae-choices">'
                f'<button class="ae-opt" onclick="aeCheck(\'{qid}\',0,{correct_idx},this)" data-correct="{correct_idx}"><span class="ae-opt-letter">✓</span>正确</button>'
                f'<button class="ae-opt" onclick="aeCheck(\'{qid}\',1,{correct_idx},this)" data-correct="{correct_idx}"><span class="ae-opt-letter">✗</span>错误</button>'
                f'</div>'
            )

        parts.append(
            f'<div class="ae-answer" id="{qid}-ans">'
            f'<span class="ae-ans-label">✅ 正确答案：</span>{ex["answer"]}'
            f'<span class="ae-ans-desc">{ex["explanation"]}</span>'
            f'</div>'
        )
        parts.append('</div>')

    return "\n".join(parts)


@router.get("/exercises/{section_id}/download")
async def download_exercises(section_id: str, subject: str = "math", grade: int = 7):
    """Generate and download a Word document of exercises."""
    ex_data = get_exercises_for(subject)
    from routes.pages import _load_sections
    sections_data = _load_sections(subject, grade)
    ex_list = ex_data.get(section_id, ex_data.get("default", []))
    if not ex_list:
        sec_name = section_id
        for s in sections_data:
            if s["id"] == section_id:
                sec_name = f"{s.get('code','')} {s['title']}"
                break
        # Still return a basic doc
        ex_list = [{"type": "choice", "question": "请先学习教材，练习题正在准备中",
                     "choices": ["好的"], "answer": 0, "explanation": ""}]

    doc = Document()

    # Page setup
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_heading('数学七年级上册 · 练习题', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sec_name = section_id
    for s in sections_data:
        if s["id"] == section_id:
            sec_name = f"{s.get('chapter','')} — {s.get('code','')} {s['title']}"
            break
    subtitle = doc.add_paragraph(sec_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('姓名：______________    日期：______________    得分：______________')
    doc.add_paragraph('')

    # Exercises
    for i, ex in enumerate(ex_list):
        q_num = f"第{i+1}题"
        q_type_map = {"choice": "【选择题】", "fill": "【填空题】"}
        q_type = q_type_map.get(ex["type"], "【题目】")

        p = doc.add_paragraph()
        run = p.add_run(f"{q_num} {q_type} {ex['question']}")
        run.bold = True
        run.font.size = Pt(11)

        if ex["type"] == "choice" and "choices" in ex:
            for j, choice in enumerate(ex["choices"]):
                doc.add_paragraph(f"    {chr(65+j)}. {choice}", style='List Bullet')

        # Answer space
        doc.add_paragraph('')
        doc.add_paragraph('答：___________________________________________________________')
        doc.add_paragraph('')

    # Answer Key page
    doc.add_page_break()
    ans_title = doc.add_heading('参考答案', level=2)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, ex in enumerate(ex_list):
        p = doc.add_paragraph()
        run = p.add_run(f"第{i+1}题答案：{ex.get('answer', '（见解析）')}")
        run.font.size = Pt(10)

        if ex.get("explanation"):
            exp_p = doc.add_paragraph()
            exp_run = exp_p.add_run(f"    解析：{ex['explanation']}")
            exp_run.font.size = Pt(9)
            exp_run.font.color.rgb = RGBColor(100, 100, 100)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Safe ASCII filename for HTTP header
    from urllib.parse import quote
    safe_name = f"math_grade7_exercises_{section_id}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_name}"},
    )


@router.get("/exercises/all", response_class=HTMLResponse)
async def all_exercises(request: Request, subject: str = "math", grade: int = 7):
    """Show all exercises from all sections (subject-aware)."""
    from routes.pages import _load_sections
    sections_data = _load_sections(subject, grade)
    ex_data = get_exercises_for(subject)
    sections = []
    for sec in sections_data:
        ex_list = ex_data.get(sec["id"], ex_data.get("default", []))
        if not ex_list:
            continue
        sections.append({"section": sec, "exercises": ex_list})

    total_ex = sum(len(s["exercises"]) for s in sections)

    return request.app.state.templates.TemplateResponse(
        "all_exercises.html",
        {"request": request, "sections": sections, "total_ex": total_ex,
         "nav_subjects": get_nav_subjects()},
    )


@router.post("/exercise/error", response_class=HTMLResponse)
async def record_error(
    section_id: str = Form(...),
    exercise_idx: int = Form(...),
    question: str = Form(...),
    correct_answer: str = Form(...),
    subject: str = Form("math"),
    db: AsyncSession = Depends(get_db),
):
    """Record a wrong answer in the error book (per-subject)."""
    stmt = select(ErrorLog).where(
        ErrorLog.subject == subject,
        ErrorLog.section_id == section_id,
        ErrorLog.exercise_idx == exercise_idx,
    )
    result = await db.execute(stmt)
    existing = result.scalar_one_or_none()

    if existing:
        existing.error_count += 1
        existing.last_error_at = datetime.now(timezone.utc)
    else:
        new_error = ErrorLog(
            subject=subject, section_id=section_id,
            exercise_idx=exercise_idx, question=question,
            correct_answer=correct_answer, error_count=1,
        )
        db.add(new_error)

    await db.commit()
    return HTMLResponse("OK")


@router.get("/error-book", response_class=HTMLResponse)
async def error_book(request: Request, subject: str = "math", db: AsyncSession = Depends(get_db)):
    """Display recorded errors for a specific subject."""
    stmt = select(ErrorLog).where(ErrorLog.subject == subject).order_by(ErrorLog.error_count.desc())
    result = await db.execute(stmt)
    errors = result.scalars().all()

    sec_names = {}
    for s in MATH_SECTIONS:
        sec_names[s["id"]] = f"{s.get('code','')} {s['title']}"

    # Enrich errors with full exercise data
    enriched = []
    for e in errors:
        ex_data = {"type": "fill", "choices": [], "explanation": ""}
        ex_subject = get_exercises_for(e.subject) if hasattr(e, 'subject') else EXERCISES
        ex_list = ex_subject.get(e.section_id, [])
        if not ex_list:
            ex_list = EXERCISES.get(e.section_id, [])
        if ex_list and e.exercise_idx < len(ex_list):
            ex_data = ex_list[e.exercise_idx]
            # For choice exercises, answer is an index — convert to display text
            if ex_data.get("type") == "choice" and "choices" in ex_data:
                try:
                    ans_idx = int(e.correct_answer)
                    if ans_idx < len(ex_data["choices"]):
                        ex_data["answer"] = ans_idx  # keep as index for template
                except (ValueError, TypeError):
                    pass
        enriched.append({
            "error": e,
            "section_name": sec_names.get(e.section_id, e.section_id),
            "exercise_data": ex_data,
        })

    total_errors = sum(e.error_count for e in errors)

    return request.app.state.templates.TemplateResponse(
        "error_book.html",
        {"request": request, "errors": enriched, "total_errors": total_errors,
         "subject": subject, "subject_name": NAMES.get(subject, subject),
         "nav_subjects": get_nav_subjects()},
    )


@router.post("/reset-all", response_class=HTMLResponse)
async def reset_all(subject: str = Form("math"), db: AsyncSession = Depends(get_db)):
    """Reset study progress and error book for a specific subject."""
    from sqlalchemy import delete
    from db.models import StudySession, ExerciseAttempt, ChatMessage, ProgressSnapshot, ErrorLog
    await db.execute(delete(ErrorLog).where(ErrorLog.subject == subject))
    # Clear sessions related to this subject's topics
    subj_sessions = (await db.execute(
        select(StudySession).where(StudySession.topic_id.like(f"{subject}%")))).scalars().all()
    for sess in subj_sessions:
        await db.execute(delete(ExerciseAttempt).where(ExerciseAttempt.session_id == sess.id))
        await db.execute(delete(ChatMessage).where(ChatMessage.session_id == sess.id))
        await db.delete(sess)
    await db.execute(delete(ProgressSnapshot))
    await db.commit()
    name = _html.escape(NAMES.get(subject, subject))
    safe_subject = _html.escape(subject)
    return HTMLResponse(f'<script>alert("✅ {name}的学习进度和错题本已重置");window.location.href="/subjects/{safe_subject}/7"</script>')


@router.get("/exercises/{section_id}", response_class=HTMLResponse)
async def get_exercises(section_id: str, subject: str = "math"):
    """Return pre-built exercises for a section as HTML."""
    html = _build_exercise_html(section_id, subject)
    js = """
<script>
function aeCheck(qid, selected, correct, el) {
    var opts = document.querySelectorAll('#' + qid + ' .ae-opt');
    opts.forEach(function(o) { o.classList.remove('ae-right','ae-wrong'); o.disabled = true; });
    if (selected === correct) { el.classList.add('ae-right'); }
    else { el.classList.add('ae-wrong'); opts[correct].classList.add('ae-right'); }
    document.getElementById(qid + '-ans').classList.add('ae-show');
}
function aeFill(qid, answer) { document.getElementById(qid + '-ans').classList.add('ae-show'); }
</script>
"""
    return HTMLResponse(html + js)
